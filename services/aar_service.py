from __future__ import annotations

from collections import OrderedDict
from dataclasses import dataclass
from datetime import datetime
from io import BytesIO
from typing import Any

from models.entities import Scenario
from services.scenario_service import ScenarioService
from services.simulation_service import SimulationService


@dataclass(frozen=True)
class ReportSection:
    title: str
    successes: list[str]
    shortfalls: list[str]
    recommendations: list[str]


class AARService:
    def __init__(
        self,
        scenario_service: ScenarioService,
        simulation_service: SimulationService,
        observer_initials: str = "BR",
    ) -> None:
        self._scenario_service = scenario_service
        self._simulation_service = simulation_service
        self._observer_initials = observer_initials

    def generate_report(self, scenario_id: str | None) -> dict[str, Any]:
        scenario = self._scenario_service.get(scenario_id)
        states = [
            self._simulation_service.build_dashboard(scenario.scenario_id, tick)
            for tick in range(scenario.duration_ticks)
        ]
        sections = self._build_sections(scenario, states)
        generated_at = self._generated_at()
        return {
            "scenario_id": scenario.scenario_id,
            "operation": scenario.operation,
            "observer": self._observer_initials,
            "generated_at": generated_at,
            "filename": f"{scenario.scenario_id}_aar.docx",
            "markdown": self._to_markdown(scenario, generated_at, sections),
            "sections": [
                {
                    "title": section.title,
                    "successes": section.successes,
                    "shortfalls": section.shortfalls,
                    "recommendations": section.recommendations,
                }
                for section in sections
            ],
        }

    def export_docx(self, scenario_id: str | None) -> tuple[BytesIO, str]:
        from docx import Document

        report = self.generate_report(scenario_id)
        buffer = BytesIO()
        document = Document()

        document.add_heading("After-Action Report (AAR)", level=0)
        document.add_paragraph(f"Operation: {report['operation']}")
        document.add_paragraph(f"Observer / Controller: {report['observer']}")
        document.add_paragraph(f"Generated: {report['generated_at']}")

        for section in report["sections"]:
            document.add_heading(section["title"], level=1)
            self._add_bullet_block(document, "Successes", section["successes"])
            self._add_bullet_block(document, "Shortfalls", section["shortfalls"])
            self._add_bullet_block(document, "Recommendations", section["recommendations"])

        document.save(buffer)
        buffer.seek(0)
        return buffer, report["filename"]

    def _build_sections(self, scenario: Scenario, states: list[dict[str, Any]]) -> list[ReportSection]:
        phase_windows = self._phase_windows(states)
        total_detections = sum(len(state["detections"]) for state in states)
        unique_targets = sorted(
            {
                detection["target_name"]
                for state in states
                for detection in state["detections"]
            }
        )
        final_state = states[-1]
        final_effects = final_state.get("active_effects", [])
        trigger_descriptions = [
            f"{trigger['condition']} -> {trigger['effect']}"
            for trigger in scenario.triggers
        ]

        sections: list[ReportSection] = [
            ReportSection(
                title="Executive Summary",
                successes=[
                    f"Mission replay completed through tick {scenario.duration_ticks - 1} for {scenario.operation}.",
                    f"ISR platforms produced {total_detections} total detection opportunities across {len(unique_targets)} unique targets."
                    if unique_targets
                    else "ISR platforms maintained coverage, but no confirmed unique targets were recorded.",
                    f"Final assessed phase was {final_state.get('current_phase', 'Baseline')}.",
                ],
                shortfalls=[
                    f"Active end-state effects were {', '.join(final_effects)}."
                    if final_effects
                    else "No terminal mission-effect flags were raised by the current rule set.",
                    f"Trigger-driven mission risks were defined as {', '.join(trigger_descriptions)}."
                    if trigger_descriptions
                    else "No explicit trigger chain was defined in the scenario payload.",
                    f"Risk score closed at {final_state['risk_summary']['score']} ({final_state['risk_summary']['label']}).",
                ],
                recommendations=[
                    "Tie mission completion criteria to explicit success and failure flags so end-state judgments are less implicit.",
                    "Preserve the current event log discipline; it gives the operator a clear replay narrative for review.",
                    "Expand rule coverage where commander intent depends on deadlines, population effects, or follow-on lift.",
                ],
            )
        ]

        for phase_name, window_states in phase_windows.items():
            phase_detections = sum(len(state["detections"]) for state in window_states)
            phase_events = [
                event["summary"]
                for state in window_states
                for event in state["event_log"]
                if event["type"] in {"uav_detection", "rule_effect"}
            ]
            sections.append(
                ReportSection(
                    title=f"Phase-by-Phase Review: {phase_name}",
                    successes=[
                        f"{phase_name} maintained {phase_detections} detection opportunities during its active window.",
                        f"Key observed events: {', '.join(dict.fromkeys(phase_events[:3]))}."
                        if phase_events
                        else f"No major rule or detection events were logged during {phase_name}.",
                        "Main Effort retained a coherent replay narrative tied to the mission timeline.",
                    ],
                    shortfalls=[
                        "Phase timing is currently driven by deterministic rule conditions rather than commander-confirmed branch decisions.",
                        "The runtime does not yet model ground losses or friction at unit echelon detail.",
                        "Phase review fidelity is bounded by the current scenario data and rule density.",
                    ],
                    recommendations=[
                        f"Add phase-specific branch conditions so {phase_name} can advance or stall based on operational outcomes.",
                        f"Attach explicit success criteria to {phase_name} for sharper AAR judgments.",
                        "Increase event specificity where the commander would want to see operational cause and effect.",
                    ],
                )
            )

        sections.append(
            ReportSection(
                title="Metrics",
                successes=[
                    f"Casualty Metric: KIA/WIA are not currently modeled, preventing false precision in the AAR.",
                    f"ROE Metric: {self._count_effects(states, {'humanitarian_penalty'})} humanitarian/ROE penalty flags were raised.",
                    f"ISR Metric: {total_detections} total detections with {len(unique_targets)} unique targets tracked.",
                    f"Logistics Metric: {self._count_effects(states, {'logistics_blockage', 'airbridge_failure'})} logistics-disruption flags were raised.",
                ],
                shortfalls=[
                    "Casualty reporting is absent from the current runtime and should be treated as a modeling gap.",
                    "ROE assessment currently depends on rule flags rather than unit-level decision traces.",
                    "ISR effectiveness is measured through detections only, not dwell quality, revisit rate, or collection gaps.",
                    "Logistics degradation is represented as flags and risk changes, not as a full sustainment model.",
                ],
                recommendations=[
                    "Add structured casualty counters before using the AAR for force employment judgments.",
                    "Track ROE breach provenance so humanitarian penalties can be attributed to specific decisions.",
                    "Add ISR persistence metrics such as time-on-station and missed detection windows.",
                    "Model logistics blockage as an effect on follow-on phases, lift availability, or fuel state.",
                ],
            )
        )

        sections.append(
            ReportSection(
                title="10% Variable: Degraded Cuban ORBAT",
                successes=[
                    "The degraded opposing ORBAT likely increased Ranger freedom of maneuver after Line of Departure by reducing armored depth and air-defense density.",
                    f"Recorded OPFOR inputs were {self._format_force_summary(scenario.opfor) or 'limited'}, which supports a faster initial penetration than a fully resourced defense.",
                    "The scenario rule set already supports deadline-based exploitation of degraded enemy posture.",
                ],
                shortfalls=[
                    "The current runtime does not calculate maneuver speed directly, so the 10% gain remains an operational estimate rather than a measured output.",
                    "Enemy degradation is inferred from scenario structure and force counts, not from a dynamic attrition model.",
                    "Follow-on friction can still erase early speed gains if logistics or runway conditions fail late.",
                ],
                recommendations=[
                    "Add explicit maneuver-tempo metrics so the estimated Ranger speed increase can be measured instead of inferred.",
                    "Represent degraded ORBAT as structured readiness values to show how enemy weakness affects tempo.",
                    "Tie enemy degradation directly to lift, seizure, and sustainment windows for cleaner operational analysis.",
                ],
            )
        )

        sections.append(
            ReportSection(
                title="82nd / Humanitarian Variable: MUVR Refugee Influx",
                successes=[
                    "The current schema can already represent ROE, refugee, and humanitarian penalties through rule effects and flags.",
                    "This gives the backend a clean path to evaluate 82nd decision-making once a refugee-flow branch is scripted.",
                    "The AAR now treats humanitarian pressure as an operational decision point rather than background noise.",
                ],
                shortfalls=[
                    "The present TRIDENT LIFT upload does not force a refugee influx at MUVR, so 82nd decision-making was not fully exercised in runtime.",
                    "Without that branch, ROE compliance and population control remain assumed rather than demonstrated.",
                    "This limits the bluntness of any claim that the 82nd balanced security, throughput, and humanitarian control effectively.",
                ],
                recommendations=[
                    "Add a dedicated MUVR refugee trigger with ROE consequences, throughput penalties, and escalation effects.",
                    "Capture 82nd decision points as explicit flags so the AAR can judge whether security or humanitarian control dominated.",
                    "Model hostile-population escalation after ROE breaches to show second-order effects on the airbridge.",
                ],
            )
        )

        sections.append(
            ReportSection(
                title="Lessons Learned",
                successes=[
                    "Document ingestion now produces scenario logic that is executable rather than purely descriptive.",
                    "The mission replay and rule engine create a review trail that an Observer / Controller can actually critique.",
                    "Scenario uploads can now be carried from document recognition to operational replay without code edits.",
                ],
                shortfalls=[
                    "Some judgments are still qualitative because ground maneuver, casualties, and humanitarian behavior are not fully modeled.",
                    "Rule quality is only as strong as the extracted or authored trigger logic inside the scenario file.",
                    "AAR accuracy still depends on disciplined scenario design and explicit mission conditions.",
                ],
                recommendations=[
                    "Continue converting commander intent into explicit triggers, flags, and metrics rather than narrative-only notes.",
                    "Favor scenarios that force tradeoffs between tempo, ROE, ISR coverage, and sustainment.",
                    "Use the AAR as a feedback loop to improve scenario logic, not just to summarize what happened.",
                ],
            )
        )

        sections.append(
            ReportSection(
                title="Recommended Logic Adjustments",
                successes=[
                    "The current schema is already flexible enough to add more detailed operational branches without refactoring the whole app.",
                    "Rule conditions and effects are sufficient for a first generation of commander-focused scenario analysis.",
                    "Exported AARs now make logic gaps visible to scenario authors.",
                ],
                shortfalls=[
                    "Ground force disposition, casualties, refugee control, and sustainment throughput remain under-modeled.",
                    "Some trigger effects still manifest as flags without enough downstream operational consequence.",
                    "The AAR must still infer certain mission outcomes because the runtime does not yet score all branches directly.",
                ],
                recommendations=[
                    "Add explicit mission-success and mission-failure scoring blocks.",
                    "Expand rule effects to include lift denial, throughput reduction, humanitarian escalation, and force readiness changes.",
                    "Add branch conditions keyed to objective completion, ROE breaches, and HVT timelines.",
                ],
            )
        )

        return sections

    def _phase_windows(self, states: list[dict[str, Any]]) -> OrderedDict[str, list[dict[str, Any]]]:
        windows: OrderedDict[str, list[dict[str, Any]]] = OrderedDict()
        for state in states:
            phase = state.get("current_phase", "Baseline")
            windows.setdefault(phase, []).append(state)
        return windows

    def _count_effects(self, states: list[dict[str, Any]], names: set[str]) -> int:
        return sum(1 for state in states for effect in state.get("active_effects", []) if effect in names)

    def _format_force_summary(self, summary: dict[str, int]) -> str:
        return ", ".join(f"{key}={value}" for key, value in summary.items())

    def _to_markdown(self, scenario: Scenario, generated_at: str, sections: list[ReportSection]) -> str:
        lines = [
            "**After-Action Report (AAR)**",
            f"- Operation: {scenario.operation}",
            f"- Theater: {scenario.theater}",
            f"- Observer / Controller: {self._observer_initials}",
            f"- Generated: {generated_at}",
            "",
        ]
        for section in sections:
            lines.extend(
                [
                    f"**{section.title}**",
                    "- Successes:",
                    *[f"  - {item}" for item in section.successes],
                    "- Shortfalls:",
                    *[f"  - {item}" for item in section.shortfalls],
                    "- Recommendations:",
                    *[f"  - {item}" for item in section.recommendations],
                    "",
                ]
            )
        return "\n".join(lines).strip()

    def _add_bullet_block(self, document: Document, title: str, items: list[str]) -> None:
        document.add_paragraph(title)
        for item in items:
            document.add_paragraph(item, style="List Bullet")

    def _generated_at(self) -> str:
        return datetime.now().astimezone().strftime("%Y-%m-%d %H:%M:%S %Z")
