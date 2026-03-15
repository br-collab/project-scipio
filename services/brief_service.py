from __future__ import annotations

from datetime import datetime
from io import BytesIO
from typing import Any

from docx import Document

from services.scenario_service import ScenarioService
from services.simulation_service import SimulationService


class BriefService:
    def __init__(
        self,
        scenario_service: ScenarioService,
        simulation_service: SimulationService,
        narrator: str = "BR",
    ) -> None:
        self._scenario_service = scenario_service
        self._simulation_service = simulation_service
        self._narrator = narrator

    def generate_brief(self, scenario_id: str | None) -> dict[str, Any]:
        scenario = self._scenario_service.get(scenario_id)
        mid_tick = min(max(1, scenario.duration_ticks // 2), scenario.duration_ticks - 1)
        state = self._simulation_service.build_dashboard(scenario.scenario_id, mid_tick)
        generated_at = datetime.now().astimezone().strftime("%Y-%m-%d %H:%M:%S %Z")
        summary = (
            f"This is how I would explain the project to a Senior SWE or CTO: I start from the user promise, "
            f"then work backward to make sure the scenario upload, replay, and map behavior visibly support that promise."
        )
        key_points = [
            f"The product claim is that {scenario.operation} becomes a runnable scenario, so I validate whether replay state, routes, and event log actually change in a credible way.",
            "I focus on the workflow boundary where planning input becomes simulation data, because that is where product trust is either created or lost.",
            "I would scope fixes narrowly: preserve the dashboard contract, keep handcrafted scenarios stable, and improve only the OPORD-to-scenario synthesis layer when uploaded scenarios feel too template-driven.",
            f"At the current replay midpoint, the dashboard shows {len(state['blue_forces'])} blue units, {len(state['red_forces'])} red threats, and {len(state['detections'])} live detections, which gives a concrete operational story rather than just parser output.",
            "The business value is translation: turning what a planner means into behavior engineering can implement and leadership can evaluate."
        ]
        return {
            "scenario_id": scenario.scenario_id,
            "operation": scenario.operation,
            "audience": "Senior SWE / CTO",
            "narrator": self._narrator,
            "generated_at": generated_at,
            "filename": f"{scenario.scenario_id}_60_second_brief.docx",
            "summary": summary,
            "key_points": key_points,
        }

    def export_docx(self, scenario_id: str | None) -> tuple[BytesIO, str]:
        brief = self.generate_brief(scenario_id)
        buffer = BytesIO()
        document = Document()

        document.add_heading("60-Second Technical Brief", level=0)
        document.add_paragraph(f"Operation: {brief['operation']}")
        document.add_paragraph(f"Audience: {brief['audience']}")
        document.add_paragraph(f"Narrator: {brief['narrator']}")
        document.add_paragraph(f"Generated: {brief['generated_at']}")
        document.add_paragraph(brief['summary'])

        document.add_heading("Talking Points", level=1)
        for point in brief['key_points']:
            document.add_paragraph(point, style='List Bullet')

        document.save(buffer)
        buffer.seek(0)
        return buffer, brief['filename']

from __future__ import annotations

from datetime import datetime
from io import BytesIO
from typing import Any

from services.scenario_service import ScenarioService
from services.simulation_service import SimulationService


class BriefService:
    def __init__(
        self,
        scenario_service: ScenarioService,
        simulation_service: SimulationService,
        narrator: str = "BR",
    ) -> None:
        self._scenario_service = scenario_service
        self._simulation_service = simulation_service
        self._narrator = narrator

    def generate_brief(self, scenario_id: str | None) -> dict[str, Any]:
        scenario = self._scenario_service.get(scenario_id)
        mid_tick = min(max(1, scenario.duration_ticks // 2), scenario.duration_ticks - 1)
        state = self._simulation_service.build_dashboard(scenario.scenario_id, mid_tick)
        generated_at = datetime.now().astimezone().strftime("%Y-%m-%d %H:%M:%S %Z")
        summary = (
            f"This is how I would explain the project to a Senior SWE or CTO: I start from the user promise, "
            f"then work backward to make sure the scenario upload, replay, and map behavior visibly support that promise."
        )
        key_points = [
            f"The product claim is that {scenario.operation} becomes a runnable scenario, so I validate whether replay state, routes, and event log actually change in a credible way.",
            "I focus on the workflow boundary where planning input becomes simulation data, because that is where product trust is either created or lost.",
            "I would scope fixes narrowly: preserve the dashboard contract, keep handcrafted scenarios stable, and improve only the OPORD-to-scenario synthesis layer when uploaded scenarios feel too template-driven.",
            f"At the current replay midpoint, the dashboard shows {len(state['blue_forces'])} blue units, {len(state['red_forces'])} red threats, and {len(state['detections'])} live detections, which gives a concrete operational story rather than just parser output.",
            "The business value is translation: turning what a planner means into behavior engineering can implement and leadership can evaluate."
        ]
        return {
            "scenario_id": scenario.scenario_id,
            "operation": scenario.operation,
            "audience": "Senior SWE / CTO",
            "narrator": self._narrator,
            "generated_at": generated_at,
            "filename": f"{scenario.scenario_id}_60_second_brief.docx",
            "summary": summary,
            "key_points": key_points,
        }

    def export_docx(self, scenario_id: str | None) -> tuple[BytesIO, str]:
        from docx import Document

        brief = self.generate_brief(scenario_id)
        buffer = BytesIO()
        document = Document()

        document.add_heading("60-Second Technical Brief", level=0)
        document.add_paragraph(f"Operation: {brief['operation']}")
        document.add_paragraph(f"Audience: {brief['audience']}")
        document.add_paragraph(f"Narrator: {brief['narrator']}")
        document.add_paragraph(f"Generated: {brief['generated_at']}")
        document.add_paragraph(brief['summary'])

        document.add_heading("Talking Points", level=1)
        for point in brief['key_points']:
            document.add_paragraph(point, style='List Bullet')

        document.save(buffer)
        buffer.seek(0)
        return buffer, brief['filename']
